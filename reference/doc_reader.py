"""
doc_reader.py  --  agent-side plain-text extraction for the FAIG RAG/file-reader surface.

Purpose: ingest an uploaded document, extract to PLAIN TEXT only, so the extracted
content can be spliced into the user turn and ride the existing FAIG-inspected path.
Base64 blobs are NOT inspectable by FAIG -- that is a separate future test, not this.

Supported: .txt .md (plain read) | .pdf (PyMuPDF) | .docx (python-docx)

This is a TEST HARNESS component. It deliberately does no sanitising of the extracted
text -- an instruction hidden inside a PDF is exactly the payload we want to survive
extraction and reach the model context, so the input scanner gets a chance to catch it.
"""

import os

# ~chars, not tokens. qwen2.5:14b context is ~32k *tokens*; ~4 chars/token English,
# but leave generous room for history + question + completion. Tune as needed.
# Private (qwen) and Public (Claude) paths diverge here -- Claude's window is larger,
# so a doc truncated for qwen may pass whole to Claude. Keep that in mind when a
# planted payload lands past the cut line.
DEFAULT_MAX_CHARS = 24000

SUPPORTED_EXTS = (".txt", ".md", ".pdf", ".docx")


class ExtractionError(Exception):
    pass


def _read_plaintext(path):
    # utf-8 with replacement so a weird byte can't kill the whole extraction
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _read_pdf(path):
    try:
        import pymupdf as fitz  # PyMuPDF >= 1.24 exposes `pymupdf`
    except ImportError:
        try:
            import fitz  # older PyMuPDF import name
        except ImportError:
            raise ExtractionError(
                "PDF support needs PyMuPDF. Install: pip install pymupdf"
            )
    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts)


def _read_docx(path):
    try:
        from docx import Document
    except ImportError:
        raise ExtractionError(
            "DOCX support needs python-docx. Install: pip install python-docx"
        )
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    # tables carry text too -- a payload could hide in a table cell
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(path, max_chars=DEFAULT_MAX_CHARS):
    """
    Return (text, meta). meta = {ext, orig_chars, truncated, returned_chars}.
    Raises ExtractionError on unsupported type or read failure.
    """
    if not os.path.isfile(path):
        raise ExtractionError(f"No such file: {path}")

    ext = os.path.splitext(path)[1].lower()
    if ext in (".txt", ".md"):
        text = _read_plaintext(path)
    elif ext == ".pdf":
        text = _read_pdf(path)
    elif ext == ".docx":
        text = _read_docx(path)
    else:
        raise ExtractionError(
            f"Unsupported type '{ext}'. Supported: {', '.join(SUPPORTED_EXTS)}"
        )

    orig = len(text)
    truncated = orig > max_chars
    if truncated:
        text = text[:max_chars] + f"\n[...truncated {orig - max_chars} chars...]"

    meta = {
        "ext": ext,
        "orig_chars": orig,
        "truncated": truncated,
        "returned_chars": len(text),
    }
    return text, meta


def build_rag_turn(question, doc_text, filename):
    """
    Splice extracted text into the user turn as explicit RAG context.

    The context/question seam is deliberate: it lets you test whether the model
    honours an instruction planted in the "context" block vs. the "question" block.
    The whole string becomes the user message -> rides the existing 4-mode routing
    -> gets scanned by FAIG's input scanner on the FAIG paths.
    """
    return (
        f'Context from uploaded document "{filename}":\n'
        f'"""\n{doc_text}\n"""\n\n'
        f'User question: {question}'
    )
